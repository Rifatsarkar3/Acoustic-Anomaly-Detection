"""Convert Revised_Manuscript_Journal_Submission.md to .docx"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH = r"e:\Yolo-Thermal\Acoustic Anomaly Detection\Revised_Manuscript_Journal_Submission.md"
DOCX_PATH = r"e:\Yolo-Thermal\Acoustic Anomaly Detection\Revised_Manuscript_Journal_Submission.docx"


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.append(tblPr)


def add_inline_formatted(paragraph, text):
    """Parse **bold** and `code` inline markup and add runs."""
    # Split on **bold** and `code` markers
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            if part:
                paragraph.add_run(part)


def parse_table_line(line):
    """Extract cells from a markdown table row."""
    line = line.strip().strip('|')
    cells = [c.strip() for c in line.split('|')]
    return cells


def is_separator_row(line):
    return bool(re.match(r'^\|[\s\-:|]+\|', line))


def process_markdown(doc, lines):
    i = 0
    while i < len(lines):
        line = lines[i]
        raw = line.rstrip()

        # Blank line
        if not raw:
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^---+$', raw):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '999999')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,4})\s+(.*)', raw)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            heading_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}
            p = doc.add_paragraph(style=heading_map.get(level, 'Heading 4'))
            add_inline_formatted(p, text)
            i += 1
            continue

        # Table: collect all consecutive table lines
        if raw.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            # Remove separator rows
            data_rows = [l for l in table_lines if not is_separator_row(l)]
            if not data_rows:
                continue
            num_cols = max(len(parse_table_line(r)) for r in data_rows)
            table = doc.add_table(rows=len(data_rows), cols=num_cols)
            table.style = 'Table Grid'
            set_table_borders(table)
            for r_idx, row_line in enumerate(data_rows):
                cells = parse_table_line(row_line)
                row = table.rows[r_idx]
                for c_idx in range(num_cols):
                    cell_text = cells[c_idx] if c_idx < len(cells) else ''
                    p = row.cells[c_idx].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if r_idx == 0:
                        run = p.add_run(cell_text)
                        run.bold = True
                    else:
                        add_inline_formatted(p, cell_text)
            doc.add_paragraph()
            continue

        # Blockquote / figure placeholder
        if raw.startswith('>'):
            text = raw.lstrip('> ').strip()
            p = doc.add_paragraph(style='Quote')
            add_inline_formatted(p, text)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.*)', raw)
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style='List Number')
            add_inline_formatted(p, text)
            i += 1
            continue

        # Bullet list (- or *)
        m = re.match(r'^[-*]\s+(.*)', raw)
        if m:
            text = m.group(1)
            p = doc.add_paragraph(style='List Bullet')
            add_inline_formatted(p, text)
            i += 1
            continue

        # Sub-bullet (indented)
        m = re.match(r'^\s{2,}[-*]\s+(.*)', raw)
        if m:
            text = m.group(1)
            p = doc.add_paragraph(style='List Bullet 2')
            add_inline_formatted(p, text)
            i += 1
            continue

        # Bold-only line (e.g. **Table X. ...**)
        m = re.match(r'^\*\*(.*)\*\*$', raw)
        if m:
            p = doc.add_paragraph()
            run = p.add_run(m.group(1))
            run.bold = True
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_inline_formatted(p, raw)
        i += 1


def main():
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Default body font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    process_markdown(doc, lines)

    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


if __name__ == '__main__':
    main()
